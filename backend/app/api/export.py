"""
Export API endpoints
Export transcripts to PDF, DOCX, TXT formats
"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import os
import logging

from app.core.database import get_collection

logger = logging.getLogger(__name__)
router = APIRouter()

EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

@router.get("/pdf/{session_id}")
async def export_pdf(session_id: str):
    """Export transcript as PDF"""
    try:
        session = await _get_session(session_id)
        
        # Create PDF
        pdf_path = os.path.join(EXPORT_DIR, f"{session_id}.pdf")
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph(f"Transcript - {session_id}", styles['Title']))
        story.append(Spacer(1, 12))
        
        # Summary
        if session.get('summary'):
            story.append(Paragraph("Summary", styles['Heading2']))
            story.append(Paragraph(session['summary'], styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Transcript
        story.append(Paragraph("Transcript", styles['Heading2']))
        for seg in session['segments']:
            text = f"{seg['speaker']} ({seg['start_time']:.1f}s): {seg['text']}"
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 6))
        
        doc.build(story)
        
        return FileResponse(
            pdf_path,
            media_type='application/pdf',
            filename=f"transcript_{session_id}.pdf"
        )
        
    except Exception as e:
        logger.error(f"PDF export error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/docx/{session_id}")
async def export_docx(session_id: str):
    """Export transcript as DOCX"""
    try:
        session = await _get_session(session_id)
        
        # Create DOCX
        doc = Document()
        doc.add_heading(f"Transcript - {session_id}", 0)
        
        # Summary
        if session.get('summary'):
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(session['summary'])
        
        # Action Items
        if session.get('action_items'):
            doc.add_heading('Action Items', level=1)
            for item in session['action_items']:
                doc.add_paragraph(f"• {item['task']}", style='List Bullet')
        
        # Transcript
        doc.add_heading('Transcript', level=1)
        for seg in session['segments']:
            text = f"{seg['speaker']} ({seg['start_time']:.1f}s): {seg['text']}"
            doc.add_paragraph(text)
        
        docx_path = os.path.join(EXPORT_DIR, f"{session_id}.docx")
        doc.save(docx_path)
        
        return FileResponse(
            docx_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"transcript_{session_id}.docx"
        )
        
    except Exception as e:
        logger.error(f"DOCX export error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/txt/{session_id}")
async def export_txt(session_id: str):
    """Export transcript as plain text"""
    try:
        session = await _get_session(session_id)
        
        # Build text content
        lines = [f"Transcript - {session_id}", "=" * 50, ""]
        
        if session.get('summary'):
            lines.append("SUMMARY")
            lines.append(session['summary'])
            lines.append("")
        
        if session.get('action_items'):
            lines.append("ACTION ITEMS")
            for item in session['action_items']:
                lines.append(f"• {item['task']}")
            lines.append("")
        
        lines.append("TRANSCRIPT")
        for seg in session['segments']:
            lines.append(f"{seg['speaker']} ({seg['start_time']:.1f}s): {seg['text']}")
        
        txt_path = os.path.join(EXPORT_DIR, f"{session_id}.txt")
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return FileResponse(
            txt_path,
            media_type='text/plain',
            filename=f"transcript_{session_id}.txt"
        )
        
    except Exception as e:
        logger.error(f"TXT export error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

async def _get_session(session_id: str):
    """Helper to get session from database"""
    collection = get_collection("transcriptions")
    session = await collection.find_one({"session_id": session_id})
    
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return session
